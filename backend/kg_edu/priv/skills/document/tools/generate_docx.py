#!/usr/bin/env python3
"""
generate_docx.py — Generate a DOCX document from Markdown content.

Uses python-docx library for professional document generation.

Input (JSON via argv[1]):
{
  "content": "# Title\\n\\nParagraph text...\\n\\n## Section\\n\\n- bullet 1\\n- bullet 2",
  "fileName": "teaching_plan",
  "outputDir": "/tmp"
}

Output (JSON to stdout):
{"filePath": "/tmp/teaching_plan.docx"}
"""
import json, sys, os, re, uuid
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_to_docx(doc, markdown_text):
    """Convert basic Markdown to docx paragraphs."""
    lines = markdown_text.split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Headers
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(14)
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(16)
        elif stripped.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(20)
        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        # Numbered list
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(text, style="List Number")
        # Normal paragraph
        else:
            p = doc.add_paragraph(stripped)


def main():
    try:
        data = json.loads(sys.argv[1])
        content = data["content"]
        file_name = data["fileName"]
        output_dir = data["outputDir"]

        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "SimSun"
        font.size = Pt(12)

        # Parse markdown
        parse_markdown_to_docx(doc, content)

        # Save
        safe_name = file_name.replace("/", "_").replace(" ", "_")[:50]
        file_path = os.path.join(output_dir, f"{safe_name}_{uuid.uuid4().hex[:8]}.docx")
        doc.save(file_path)

        print(json.dumps({"filePath": file_path}))

    except Exception as e:
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
