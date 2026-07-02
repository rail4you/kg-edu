#!/usr/bin/env python3
"""
智课云枢 项目参数文档生成脚本
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=11, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_bullet(text, level=0, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * (level + 1))
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = '微软雅黑'
        run.font.size = Pt(11)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.name = '微软雅黑'
        run.font.size = Pt(10)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if header:
            run.bold = True
            set_cell_shading(cell, "D9E2F3")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return row

def make_table(headers, data, col_widths=None):
    table = doc.add_table(rows=0, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    add_table_row(table, headers, header=True)
    for row_data in data:
        add_table_row(table, row_data)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('智课云枢 智慧教学系统')
run.font.name = '微软雅黑'
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('项目建设技术参数说明书')
run.font.name = '微软雅黑'
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run('（基于知识图谱与人工智能技术的智慧课程升级方案）')
run.font.name = '微软雅黑'
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()
doc.add_paragraph()

ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver_p.add_run('版本：V2.0')
run.font.name = '微软雅黑'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

# ════════════════════════════════════════════════
# 目录页（手动）
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('目 录', level=1)
toc_items = [
    '一、项目建设背景与目标',
    '二、系统总体架构',
    '  2.1 技术架构',
    '  2.2 功能架构',
    '  2.3 部署架构',
    '三、教师建课空间',
    '  3.1 课程创建与管理',
    '  3.2 课程教学互动',
    '  3.3 考核与证书',
    '  3.4 题库管理',
    '  3.5 作业管理',
    '  3.6 测验与考试管理',
    '  3.7 数据统计与教学大屏',
    '  3.8 消息通知',
    '四、知识图谱核心功能',
    '  4.1 图谱创建与导入导出',
    '  4.2 图谱编辑（三种模式）',
    '  4.3 图谱协作',
    '  4.4 图谱统计',
    '  4.5 学情图谱',
    '  4.6 职业图谱（岗赛证）',
    '五、课程智能体（AI Agent）',
    '  5.1 智能体建设与管理',
    '  5.2 语料管理与审核',
    '  5.3 问答对管理',
    '六、AI 工具集',
    '  6.1 AI 批阅',
    '  6.2 AI 课程教学报告',
    '  6.3 AI 学习助手',
    '  6.4 AI 视频理解',
    '  6.5 AI 文档解析',
    '  6.6 AI 错题解析',
    '  6.7 AI 学习报告',
    '  6.8 AI 练习题生成',
    '  6.9 全平台内容安全（绿网）',
    '七、微专业功能',
    '  7.1 微专业管理',
    '  7.2 微专业课程体系',
    '  7.3 学生报名与审批',
    '  7.4 微专业学习门户',
    '  7.5 运营分析与数据看板',
    '八、分组教学与课堂互动',
    '  8.1 分组管理',
    '  8.2 分组任务',
    '九、个性化学习空间',
    '  9.1 课程学习',
    '  9.2 学习计划',
    '  9.3 收藏夹与证书',
    '  9.4 认知目标看板',
    '  9.5 学习推荐',
    '  9.6 邮件问答',
    '十、专业分析与人才培养',
    '  10.1 专业管理',
    '  10.2 岗位管理与 AI 分析',
    '  10.3 能力图谱构建',
    '  10.4 AI 课程体系设计',
    '  10.5 分析报告',
    '十一、学生学情画像',
    '  11.1 总览卡片',
    '  11.2 知识掌握与能力雷达',
    '  11.3 学习行为与趋势',
    '  11.4 薄弱知识点分析',
    '十二、系统管理',
    '  12.1 虚拟教研室',
    '  12.2 AI 工作台',
    '  12.3 系统设置',
    '十三、部署与售后服务要求',
]
for item in toc_items:
    add_para(item, size=12, space_after=2)

# ════════════════════════════════════════════════
# 一、项目建设背景与目标
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('一、项目建设背景与目标', level=1)

add_para('为积极落实国家教育数字化战略行动，推动人工智能赋能教育变革转型，实现信息技术与教育教学深度融合，提高课程建设质量，助力专业数智化转型，提升人才培养成效，拟对现有在线精品课程进行"人工智能＋"智能升级，打造"智课云枢"智慧教学系统。')
add_para('"智课云枢"利用大模型、知识图谱等人工智能技术与职业教育/高等教育深度融合，为学校打造智能化、创新性、实用性的智慧课程平台。系统以"知识图谱"为核心驱动，以"AI 智能体"为智慧引擎，以"微专业"为特色培养路径，构建覆盖教、学、管、评、研全链条的智慧教学生态系统。')

# ════════════════════════════════════════════════
# 二、系统总体架构
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('二、系统总体架构', level=1)

add_heading_styled('2.1 技术架构', level=2)
add_para('系统采用前后端分离的云原生架构设计：')
make_table(
    ['层次', '技术栈', '说明'],
    [
        ['前端展示层', 'React 18 + TypeScript + Vite + Ant Design Pro', 'SPA 单页应用，响应式设计，支持 PC/移动端'],
        ['API 网关层', 'Vite Proxy（开发）/ Nginx（生产）', '请求路由、反向代理、路径重写'],
        ['后端服务层', 'Elixir + Phoenix + Ash Framework', '声明式领域建模，REST + RPC 双模式 API'],
        ['AI 引擎层', '大语言模型 + 课程智能体 + AI 工具集', '基于LLM的智能批阅、生成、对话、分析'],
        ['数据持久层', 'PostgreSQL（租户级多租户）', '每个租户独立 schema，数据隔离'],
        ['对象存储', '阿里云 OSS（兼容 S3 协议）', '静态资源、文件上传、视频存储'],
    ],
    col_widths=[3, 6, 7.5]
)

add_heading_styled('2.2 功能架构', level=2)
add_para('系统采用双模块 + 多中心的总体功能架构：')
add_bullet('智慧课程模块：课程管理、知识图谱、AI工具集、教学互动、考核评价、统计分析')
add_bullet('微专业模块：微专业管理、独立课程体系、学生报名审批、运营分析')
add_bullet('专业分析中心：专业管理、岗位分析、能力图谱、课程体系设计、分析报告')
add_bullet('AI 工作台：AI智能体、AI生成试题、AI文件管理、AI命令中心')

add_heading_styled('2.3 部署架构', level=2)
make_table(
    ['环境', '组件', '端口', '说明'],
    [
        ['开发环境', 'Vite Dev Server', '8081', '前端热更新开发服务器，含 proxy 转发'],
        ['开发环境', 'Phoenix Backend', '4000', '后端 API + 静态文件服务'],
        ['开发环境', 'PostgreSQL', '5433', '开发数据库'],
        ['生产环境', 'Phoenix Backend', '4000', '统一容器，服务 SPA + REST API + LiveView'],
        ['生产环境', 'PostgreSQL', '5432', '生产数据库'],
        ['生产环境', 'Nginx（可选）', '80', '反向代理 + SSL 终止'],
    ],
    col_widths=[2.5, 4, 2, 8]
)

# ════════════════════════════════════════════════
# 三、教师建课空间
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('三、教师建课空间', level=1)

add_heading_styled('3.1 课程创建与管理', level=2)
add_bullet('新建课程：支持从零创建课程，或从原有平台同步现有在线精品课程与资源库课程至平台')
add_bullet('课程信息编辑：课程封面、课程介绍视频、课程简介、关键词、教材信息等')
add_bullet('章节结构管理：支持章/节/讲三级结构，支持拖拽排序、删除、重命名等操作')
add_bullet('教学团队维护：添加/删除教学团队成员，展示团队成员信息，支持修改简介等个人信息')
add_bullet('学员管理：按姓名、学校、选课时间等多维搜索查询学员，支持导出课程学员名单')
add_bullet('学科分类管理：课程按学科分类组织，便于检索与统计')
add_bullet('教师课程分配：支持设置主讲教师和助理教师角色，灵活管理授课权限')
add_bullet('学生选课管理：支持选课审核和手动添加/移除选课学生')

add_heading_styled('3.2 课程教学互动', level=2)
add_bullet('课程公告：新建、编辑、删除公告，支持关键词查询')
add_bullet('主题讨论：创建/编辑/删除讨论主题，支持点赞、回复、设置精华帖')
add_bullet('课程答疑：学生提问 → 教师回答（支持人工回答和 AI 智能回答两种模式）')
add_bullet('课程笔记：按章节筛选查询学生笔记，支持点赞互动')
add_bullet('课程评价：教师可查询、管理课程评价内容')

add_heading_styled('3.3 考核与证书', level=2)
add_bullet('成绩区间设置：拖拽式设置不合格/合格/优秀三级的成绩区间范围')
add_bullet('权重组成：学生成绩由资源完成率、考试、在线作业、测验、主题讨论五项组成，教师可自定义权重')
add_bullet('通过标准：支持设置课程通过分数线')
add_bullet('明细配置：考试、在线作业、测验、主题讨论支持教师配置参与考核的具体内容')
add_bullet('证书设置：支持不开启证书、合格证书、合格+优秀证书三种模式，可设置对应分数线')

add_heading_styled('3.4 题库管理', level=2)
add_bullet('多源导入：支持从原课程平台的其他课期批量导入题目')
add_bullet('多维筛选：按题型、题目难度、题目来源进行筛选，支持关键词搜索')
add_bullet('题库维护：删除题目、修改题目难度、批量导出题目')
add_bullet('题型支持：单选题、多选题、判断题、填空题、简答题、问答题、综合题等多种题型')

add_heading_styled('3.5 作业管理', level=2)
add_para('（1）在线作业：', bold=True)
add_bullet('题目来源为课程题库，支持设置答题时长、完成时间、作答次数', level=1)
add_bullet('防作弊设置：题目乱序、选项乱序', level=1)
add_bullet('多选题支持多种计分方式', level=1)
add_bullet('成绩认定方式、答案公布时间、成绩公布时间灵活设置', level=1)
add_bullet('随机出题和手动出题两种模式：随机出题在确定题型和数目后随机组卷，保证学生收到差异化试卷', level=1)

add_para('（2）附件作业：', bold=True)
add_bullet('支持创建/编辑/删除附件作业，可编辑作业名称、答题时长、完成时间、答题次数')
add_bullet('支持上传作业说明和参考答案附件')
add_bullet('支持设置答案公布时间和成绩公布时间')

add_para('（3）批阅方式：', bold=True)
add_bullet('客观题系统自动批阅')
add_bullet('主观题支持 AI 智能批阅（含评分和评语）')
add_bullet('教师可在线查看学生作业作答情况')

add_heading_styled('3.6 测验与考试管理', level=2)
add_bullet('支持创建/编辑/删除测验和考试，功能与在线作业基本一致')
add_bullet('支持多种考试类型配置')
add_bullet('批改功能：支持教师在线批改试卷')
add_bullet('成绩分析：提供考试统计和分析功能')
add_bullet('支持按时间范围筛选考试')

add_heading_styled('3.7 数据统计与教学大屏', level=2)
add_bullet('成果统计：本期选课人数、学员所属单位、互动次数、日志总数、教学团队统计')
add_bullet('知识图谱建设数据展示：图谱建设统计、专项图谱数据')
add_bullet('教学资源统计：课程、章节、知识点、视频、文档数量及总时长')
add_bullet('教学活动统计：作业、练习题、AI 生成练习等')
add_bullet('AI 功能使用统计：AI命令使用次数、AI练习生成数量')
add_bullet('教学大屏：学习数据统计展示、教学活动数据统计展示、成绩统计展示')
add_bullet('课程大数据分析：学习趋势、学习进度、成绩趋势、考评完成度、观看视频时长等')

add_heading_styled('3.8 消息通知', level=2)
add_bullet('查看已发送信息及收到的信息')
add_bullet('支持一键已读、邮件状态跟踪')

# ════════════════════════════════════════════════
# 四、知识图谱核心功能
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('四、知识图谱核心功能', level=1)

add_para('知识图谱是本系统的核心驱动引擎，依托平台的知识图谱构建功能，支持建设课程专属知识图谱，并在教师建课后台、前端门户实时展示。系统提供从图谱创建、编辑、协作到多维可视化展示的完整工具链。')

add_heading_styled('4.1 图谱创建与导入导出', level=2)
add_bullet('创建课程后自动创建对应图谱；导入课程可基于现有课程结构自动生成与章节对应的图谱')
add_bullet('支持依据导入模板格式一键导入，自动生成图谱')
add_bullet('支持引用课程上一期的图谱，自动将上一期的节点资源、节点标签等同步到新一期')
add_bullet('课程名称、封面、专业大类信息自动同步至对应图谱')
add_bullet('支持将图谱内容一键导出')

add_heading_styled('4.2 图谱编辑（三种模式）', level=2)
add_para('系统提供预览和编辑两种模式，满足老师在不同场景的使用需求。提供三种图谱编辑方式：', bold=False)

make_table(
    ['编辑模式', '特点', '适用场景'],
    [
        ['知识技能导图', '自由分布布局，节点可任意拖拽', '宏观知识结构展示'],
        ['知识技能树', '树形结构，支持拖拽调整章/节/讲/资源的顺序', '与课程结构一致的精细化编排'],
        ['知识技能图谱', '网络关系布局，展示节点间复杂关联', '知识点交叉关联展示'],
    ],
    col_widths=[3.5, 5.5, 5]
)

add_para('')
add_para('核心编辑能力：', bold=True)
add_bullet('根据布鲁姆认知学习领域目标（记忆、理解、应用、分析、评价、创造）搭建知识图谱标签体系')
add_bullet('完善图谱节点基本信息：节点名称、节点描述、节点目标、节点属性')
add_bullet('标记节点属性：是否关联岗位、是否关联证书、是否为赛点、是否为课程思政点')
add_bullet('支持设置和自定义添加节点标签')
add_bullet('AI 助写功能：智能生成图谱节点描述、知识节点目标')
add_bullet('节点关系编辑：支持添加包含、并列、先后关系')

add_heading_styled('知识点关联', level=3)
add_bullet('支持将图谱节点与统一资源中心、专业资源库、本地的资源相关联，建立知识点与资源的绑定关系')
add_bullet('支持将图谱节点与课程的作业、测验、考试相关联，建立知识点与考核的绑定关系')
add_bullet('知识关系管理：支持设置前置知识、后置知识、依赖关系等多种关联类型')

add_heading_styled('4.3 图谱协作', level=2)
add_bullet('教学团队成员参与图谱协作，课程负责人可分配协作节点、设置任务截止时间')
add_bullet('课程负责人拥有全部课程节点的编辑权限')
add_bullet('支持修改、撤销协作任务')
add_bullet('团队成员编辑完成后提交至课程主持人审核')

add_heading_styled('4.4 图谱统计', level=2)
add_bullet('提供各级节点数、关联岗位数、关联证书数、关联赛项数的统计')
add_bullet('关联资源数、关联作业测验考试数、团队成员数、点击次数、访问次数')

add_heading_styled('4.5 学情图谱（学生端）', level=2)
add_para('学情图谱以图谱形式展示学生的学习进度、测评分数和薄弱知识点，帮助学生直观了解个人学习情况：')

add_para('（1）进度图谱', bold=True)
add_bullet('以图谱形式直观呈现课程每个节点的学习进度')
add_bullet('查看每个节点下的学习资源列表及各资源的学习进度')
add_bullet('点击弹窗可直接跳转资源进行学习')

add_para('（2）测评图谱', bold=True)
add_bullet('以图谱形式呈现每个节点的作业、测验、考试列表及作答成绩')
add_bullet('弹窗中可直接进入试卷作答')

add_para('（3）薄弱知识点图谱', bold=True)
add_bullet('以图谱形式展示每个节点的资源、作业、测验、考试的整体掌握程度')
add_bullet('同步展示掌握程度、学习进度及测评得分率相关数据')

add_heading_styled('4.6 职业图谱（岗赛证）', level=2)
add_para('职业图谱呈现课程内容与岗位、竞赛、证书之间的关联关系，辅助学生根据职业发展需求进行学习：')

add_para('（1）岗位图谱', bold=True)
add_bullet('以图谱形式直观呈现课程节点与岗位信息的关联')
add_bullet('展示岗位的工作领域、工作任务、技能描述与课程节点的关联关系')
add_bullet('支持查看节点设置的岗位、证书等信息及具体资源')

add_para('（2）比赛图谱', bold=True)
add_bullet('以图谱形式呈现课程节点与比赛及赛项信息的关联')
add_bullet('展示各赛项、赛点与课程节点的关联关系')

add_para('（3）证书图谱', bold=True)
add_bullet('以图谱形式呈现课程节点与证书信息的关联')
add_bullet('展示证书考点与课程节点的对应关系')

add_para('')
add_para('可视化视图模式：', bold=True)
add_bullet('支持树图、环图、层次结构图、问题图、思政图谱、能力图谱等多种视图切换')
add_bullet('知识点搜索定位、节点缩放拖拽')
add_bullet('课程图谱支持自由分布和有序呈现两种展示方式')

# ════════════════════════════════════════════════
# 五、课程智能体
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('五、课程智能体（AI Agent）', level=1)

add_para('基于大语言模型核心能力，训练解析基于本课程的课程级智能体，提供课程级别的智能问答助手。智能体负责人、训练人员、审核人员三位一体协同管理。')

add_heading_styled('5.1 智能体建设与管理', level=2)
add_bullet('智能体负责人查看已分配的智能体项目信息列表')
add_bullet('为项目添加多个训练人员和审核人员')
add_bullet('智能体对话在讨论区和答疑区自动接入，智能解答学生提问')

add_heading_styled('5.2 语料管理与审核', level=2)
add_bullet('语料新增/修改/删除/审核/查看')
add_bullet('支持解析多模态资源：视频、音频、文档（PPT、Word、Excel、PDF等）')
add_bullet('查看语料资源信息：资源名称、格式类型、资源容量、数据来源、上传时间、上传人、所属单位、审核人、状态、操作等')
add_bullet('训练人员可新增、添加语料，查看语料详情及状态')
add_bullet('审核人员查看待审核/已审核的语料资源，执行审核操作')

add_heading_styled('5.3 问答对管理', level=2)
add_bullet('课程问答对管理：支持从题库中选择题目生成问答对，或手动创建问答对')
add_bullet('问答对新增/修改/删除/审核/查看')
add_bullet('训练人员查看已提交、待审核、通过、未通过的问答对列表')
add_bullet('审核人员可对问答对详情数据进行查看及审核')

# ════════════════════════════════════════════════
# 六、AI 工具集
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('六、AI 工具集', level=1)

add_para('依托大模型与智能体技术，系统提供覆盖教学全场景的 AI 工具矩阵：')

add_heading_styled('6.1 AI 批阅', level=2)
add_bullet('AI自动批阅作业与考试，结合评分标准剖析答案内容，自动打分并提供智能批语')
add_bullet('总结性评价：自动分析学生回答，生成评语和评分')
add_bullet('逐句评价：对内容提供细致的逐句评价结果')
add_bullet('多维度评价：针对学习态度、认知能力、表达能力等方面进行评价')

add_heading_styled('6.2 AI 课程教学报告', level=2)
add_bullet('基于教学过程数据智能生成课程教学报告')
add_bullet('评价指标：学习用户量、成绩分布、进度分布、课程建设内容')
add_bullet('关键数据洞察：表现最突出的学生、最受欢迎的问题、整体分数最好的考试、最需关注的知识点')
add_bullet('知识点掌握程度评价、学生考测评排行榜')
add_bullet('课程风格分析：以视频资源为主、题库丰富、讨论气氛热烈等')
add_bullet('教学详情分析：参与情况、课程设计情况、智能工具使用情况、知识图谱建设情况、考测评情况')
add_bullet('重点内容汇聚：登录时间分布、学生讨论关键词、视频解析及错题关键词、热门对话问题TOP5、班级薄弱知识图谱')
add_bullet('智能生成课程建设参考建议')

add_heading_styled('6.3 AI 学习助手', level=2)
add_bullet('课程学习助手：配套展示在课程学习页面，支持学生随时调用，进行基于课程内容的过程性问答')
add_bullet('讨论支持：在讨论区自动解答学生发起的讨论问题')
add_bullet('答疑支持：在答疑区自动解答学生提问')

add_heading_styled('6.4 AI 视频理解', level=2)
add_bullet('AI 视频速览：多模态技术对视频切片，梳理知识点，支持按知识点跳转对应视频片段')
add_bullet('AI 视频字幕及实时翻译：自动生成视频字幕，支持中英双语')
add_bullet('AI 视频总结：基于视频内容自动提取生成概要总结')
add_bullet('AI 热词解析：自动生成热词词云，点击词云即可基于词云对话')

add_heading_styled('6.5 AI 文档解析', level=2)
add_bullet('支持 Word、PDF 等文档格式上传，基于文档进行问答对话')
add_bullet('基于文档的分层级总结，重点内容特殊标记')

add_heading_styled('6.6 AI 错题解析', level=2)
add_bullet('学生遇到错题且无解析时，AI 自动智能解析错因')
add_bullet('同步提供相关知识点资源展示，帮助针对性巩固')

add_heading_styled('6.7 AI 学习报告', level=2)
add_bullet('根据学生参与度、成绩、学习风格和偏好、学习路径和薄弱知识点分析生成个性化 AI 学习报告')
add_bullet('参与度分析：登录频率、学习时长、资源学习、讨论答疑、测验完成情况 → 智能评价')
add_bullet('学习习惯分析：材料类型偏好、学习时间分布、成绩趋势、难点掌握 → 学习建议')
add_bullet('学习详情展示：参与度（登录/资源/作业/讨论）、智能工具使用（错题解析/视频解析/智能体对话/图谱使用）、学习成果智能评价')

add_heading_styled('6.8 AI 练习题生成', level=2)
add_bullet('利用 AI 技术自动生成练习题目，支持多种题型')
add_bullet('基于课程内容智能出题，支持按课程和题型筛选')

add_heading_styled('6.9 全平台内容安全（绿网）', level=2)
add_bullet('全平台智能回复接入"绿网"功能')
add_bullet('运用语义分析与内容识别技术，构建多层级内容过滤机制')
add_bullet('对所有 AI 生成回复进行实时精准审查')
add_bullet('快速识别不良或敏感信息：低俗内容、错误价值观导向言论、不实知识信息等')

# ════════════════════════════════════════════════
# 七、微专业功能
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('七、微专业功能', level=1)

add_para('微专业模块是本系统的特色人才培养功能，提供从微专业创建、独立课程体系构建、学生报名审批到学习跟踪的完整闭环。微专业与智慧课程模块数据独立，但支持数据导入复用。')

add_heading_styled('7.1 微专业管理', level=2)
add_bullet('教师端独立微专业模块，左侧导航专属功能菜单')
add_bullet('微专业创建/编辑/发布/归档全生命周期管理')
add_bullet('发布状态管理：草稿、已发布（可报名）、已归档')
add_bullet('从智慧课程模块导入课程资源（视频、习题、资源）')

add_heading_styled('7.2 微专业课程体系', level=2)
add_para('微专业拥有独立课程体系（与智慧课程模块数据隔离但结构一致）：')
make_table(
    ['数据模型', '说明', '与智慧课程关系'],
    [
        ['MicroMajorCourse', '微专业独立课程实体', '与智慧课程同构但数据独立'],
        ['MicroMajorChapter', '课程章节（支持多级自嵌套）', '树形结构，parent_chapter_id 自引用'],
        ['MicroMajorVideo', '微专业视频资源', '可从智慧课程导入'],
        ['MicroMajorExercise', '微专业习题', '可从智慧课程导入，支持多种题型'],
        ['MicroMajorResource', '微专业文件资源', '可从智慧课程导入'],
    ],
    col_widths=[4, 6, 4.5]
)
add_para('')
add_bullet('课程管理：CRUD + 发布/下架状态控制')
add_bullet('章节管理：树形结构 + 拖拽排序')
add_bullet('视频管理：上传 + 列表管理')
add_bullet('习题管理：CRUD + 题型支持（单选/多选/判断/填空/问答等）')
add_bullet('资源管理：文件上传 + 列表管理')
add_bullet('导入功能：从智慧课程的多选/搜索表格界面批量导入视频、习题、资源')

add_heading_styled('7.3 学生报名与审批', level=2)
add_para('微专业支持学生自主报名 + 教师审批的完整流程：', bold=True)

make_table(
    ['状态', '说明', '触发动作'],
    [
        ['pending（待审核）', '学生已报名，等待教师审批', '学生报名或重新报名'],
        ['active（已通过）', '审批通过，学生可正常学习', '教师批准 / 教师手动分配'],
        ['rejected（已拒绝）', '教师拒绝报名', '教师拒绝时填写拒绝原因'],
        ['completed（已完成）', '学生完成微专业学习', '系统自动或手动设置'],
        ['removed（已移除）', '学生被移除', '教师手动移除'],
    ],
    col_widths=[3, 5, 6]
)
add_para('')
add_bullet('学生端：浏览公开微专业列表 → 点击"报名申请" → 查看审核状态（审核中/已通过/已拒绝）')
add_bullet('被拒后支持"重新报名"，清空拒绝原因重新进入待审核')
add_bullet('教师端审批管理：待审核/已通过/已拒绝三个 Tab')
add_bullet('审批操作：批准（可填写备注）或 拒绝（需填写拒绝原因）')

add_heading_styled('7.4 微专业学习门户', level=2)
add_bullet('学生端微专业首页：卡片网格布局，展示封面、名称、学习周期、进度条')
add_bullet('微专业详情页：课程列表（标题/描述/学分/学时）')
add_bullet('课程学习页面：四个 Tab—课程介绍 + 章节树 / 视频学习（播放器+章节导航）/ 习题练习 / 资源下载')
add_bullet('学习行为记录：ActivityLog 记录视频观看、习题提交、文件下载等操作')

add_heading_styled('7.5 运营分析与数据看板', level=2)
add_bullet('教师端运营分析面板：微专业 → 课程 → 学生三层维度')
add_bullet('数据指标：选课人数、视频观看数、习题完成率、资源下载数')
add_bullet('学生学习时间轴：按日期分组的 ActivityLog 动态展示')

# ════════════════════════════════════════════════
# 八、分组教学与课堂互动
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('八、分组教学与课堂互动', level=1)

add_para('教师可在课堂互动中创建学生分组、布置分组任务，通过二维码分发学生参与，教师端实时跟踪完成进度。')

add_heading_styled('8.1 分组管理', level=2)
add_bullet('创建小组：名称/描述 + 成员管理')
add_bullet('随机分组：一键自动将课程学生随机分成 N 组')
add_bullet('小组维护：添加/移除成员')

add_heading_styled('8.2 分组任务', level=2)
add_bullet('创建任务：标题/类型/描述/截止日期')
add_bullet('任务类型：提交任务（文本+文件上传）、讨论任务、投票调查、文件上传')
add_bullet('发布任务：自动生成二维码弹窗 + 分享链接')
add_bullet('进度跟踪：状态管理（草稿/进行中/已结束）')
add_bullet('学生扫码进入 → 查看任务 → 提交内容 → 完成标记')

# ════════════════════════════════════════════════
# 九、个性化学习空间
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('九、个性化学习空间', level=1)

add_heading_styled('9.1 课程学习', level=2)
add_bullet('内容学习：视频、图片、文档、富媒体、音频等课程资源的在线学习，记录学习过程数据')
add_bullet('课程公告查看')
add_bullet('作业/测验/考试：在线答题、提交、查看成绩')
add_bullet('个人成绩：综合成绩、分项得分')
add_bullet('课程讨论：话题评论、回复、点赞')
add_bullet('课程答疑：答疑区发起问题，教师或 AI 解答')
add_bullet('课程笔记：学习时创建笔记')

add_heading_styled('9.2 学习计划', level=2)
add_bullet('学生制定学习计划，系统根据计划进行消息提示与督促')
add_bullet('实时更新展示学习计划完成程度')

add_heading_styled('9.3 收藏夹与证书', level=2)
add_bullet('收藏课程，收藏夹查看')
add_bullet('个人证书：申请/查看/下载已通过的课程证书')

add_heading_styled('9.4 认知目标看板', level=2)
add_bullet('展示课程认知目标学习看板，覆盖记忆/理解/应用/分析/评价/创造六个层级')
add_bullet('知识点认知层级分布可视化展示')
add_bullet('树形知识点列表 + 各认知目标的设置情况')

add_heading_styled('9.5 学习推荐', level=2)
add_bullet('基于协同过滤的个性化学习推荐')
add_bullet('推荐维度：同学学习同类资源、本课程热门资源、知识点掌握度推荐')
add_bullet('状态跟踪：待学习 / 进行中 / 已完成')

add_heading_styled('9.6 邮件问答', level=2)
add_bullet('学生向教师发送提问邮件')
add_bullet('教师回复后学生可查看历史问答记录')
add_bullet('发送状态跟踪：已发送/发送失败')

# ════════════════════════════════════════════════
# 十、专业分析与人才培养
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('十、专业分析与人才培养', level=1)

add_para('专业分析模块基于录入的专业信息，调用 AI 服务开展岗位分析、能力分析，构建专业能力素质图谱，AI 辅助人才培养设计。')

add_heading_styled('10.1 专业管理', level=2)
add_bullet('专业增删改查：名称、代码、学院、学位类型、学制等')
add_bullet('专业详情总览页，提供四个子模块入口卡片')

add_heading_styled('10.2 岗位管理与 AI 分析', level=2)
add_bullet('岗位管理：添加岗位名称/描述/要求/薪资')
add_bullet('AI 岗位分析：自动分析核心技能、发展前景、能力要求')
add_bullet('岗位详情查看 + AI 分析结果展示')

add_heading_styled('10.3 能力图谱构建', level=2)
add_bullet('手动添加能力节点（名称/类别/级别/权重）')
add_bullet('能力类别：专业能力、通用能力、实践能力')
add_bullet('支持树形层级结构')

add_heading_styled('10.4 AI 课程体系设计', level=2)
add_bullet('AI 基于专业信息和能力数据自动生成课程体系方案')
add_bullet('课程数据预览（JSON格式）')
add_bullet('发布方案管理')

add_heading_styled('10.5 分析报告', level=2)
add_bullet('AI 自动生成四种类型分析报告：')
add_bullet('岗位分析报告', level=1)
add_bullet('能力图谱报告', level=1)
add_bullet('课程体系报告', level=1)
add_bullet('综合分析报告', level=1)
add_bullet('报告格式：Markdown 在线预览')

# ════════════════════════════════════════════════
# 十一、学生学情画像
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('十一、学生学情画像', level=1)

add_para('整合学生多维度学习数据，构建动态数字画像，立体呈现学员知识掌握、学习行为与能力素养。')

make_table(
    ['画像区域', '图表类型', '展示内容'],
    [
        ['总览卡片', '统计卡片（4个）', '知识掌握度、学习活跃度、考试平均分、出勤率'],
        ['知识掌握', '雷达图', '按主能力维度展示掌握度，与班级平均对比'],
        ['学习行为', '饼图', '视频学习/资料阅读/练习答题/作业提交占比'],
        ['学习趋势', '折线图', '按周聚合的学习活动趋势'],
        ['能力维度', '柱状图', '各子能力得分与班级平均对比'],
        ['薄弱知识点', '表格', '掌握度低于阈值的知识点列表'],
        ['班级薄弱分析', '列表', '全班最薄弱的 TOP N 知识点'],
    ],
    col_widths=[3, 3.5, 7.5]
)

add_para('')
add_bullet('教师可在课程和学生的两级选择器下查看任意学生的完整画像')
add_bullet('数据来源：知识点掌握度、活动日志、考试成绩、出勤记录、能力维度')

# ════════════════════════════════════════════════
# 十二、系统管理
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('十二、系统管理', level=1)

add_heading_styled('12.1 虚拟教研室', level=2)
add_bullet('教师团队信息展示：卡片形式展示所有教师信息')
add_bullet('每位教师：头像、姓名、职称、联系方式、简介')
add_bullet('展示课程创建数量和分配课程数量')
add_bullet('教师简介编辑')

add_heading_styled('12.2 AI 工作台', level=2)

add_para('（1）AI 智能体助手', bold=True)
add_bullet('智能对话服务，帮助教师解答教学相关问题')
add_bullet('会话历史列表 + 主聊天界面')
add_bullet('新建/删除会话，对话反馈（点赞/点踩/复制）')

add_para('（2）AI 生成试题', bold=True)
add_bullet('自动生成练习题，按课程和题型筛选')
add_bullet('题目类型选择 + 生成历史')

add_para('（3）AI 文件管理', bold=True)
add_bullet('管理 AI 生成的文件内容')
add_bullet('卡片/表格双视图切换 + 搜索')

add_para('（4）AI 命令中心', bold=True)
add_bullet('管理 AI 命令模板，定制个性化助手指令')
add_bullet('创建/编辑/删除命令模板')

add_heading_styled('12.3 系统设置', level=2)
add_bullet('界面设置：语言选择（简体中文）')
add_bullet('AI 设置：AI Key 配置、模型选择')
add_bullet('账户信息：用户名、邮箱、密码修改')

# ════════════════════════════════════════════════
# 十三、部署与售后服务
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('十三、部署与售后服务要求', level=1)

add_bullet('部署时限：中标人应及时完成平台部署，并对采购方使用人员进行操作及日常维护培训')
add_bullet('售后服务：指派专人通过企业微信、在线客服工具提供 7×24 小时在线服务')
add_bullet('响应时效：学校提出问题，2 小时内给予响应')
add_bullet('云计算架构：学校不需要准备本地存储设备')
add_bullet('服务期内：由中标人承担云存储空间及带宽加速等云端软硬件环境租用费用')
add_bullet('包含：云服务器、带宽、CDN 加速等')
add_bullet('服务期满：平台升级、运维保障及云存储空间服务费用由双方届时另行协商')

# ════════════════════════════════════════════════
# 附录：技术参数汇总表
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('附录：技术参数汇总表', level=1)

make_table(
    ['技术指标', '参数/说明'],
    [
        ['前端技术栈', 'React 18 + TypeScript + Vite + Ant Design Pro'],
        ['后端技术栈', 'Elixir + Phoenix + Ash Framework'],
        ['API 模式', 'REST API + Typescript RPC 双模式'],
        ['数据库', 'PostgreSQL，多租户 Schema 隔离'],
        ['认证方式', 'JWT + AshAuthentication'],
        ['AI 模型接入', '支持多种大语言模型接入（LLM）'],
        ['对象存储', '兼容 S3 协议（阿里云 OSS 等）'],
        ['容器化', 'Docker 部署（db + backend 双容器）'],
        ['多租户', '基于 Schema 的租户级数据隔离'],
        ['知识图谱模式', '导图/树图/网络图三种编辑模式'],
        ['可视化视图', '树图/环图/层次图/问题图/思政图/能力图'],
        ['微专业数据模型', '独立课程体系（5 张核心表）'],
        ['学习行为追踪', 'ActivityLog + StudentKnowledgeMastery'],
        ['内容安全', 'AI 生成内容实时审查过滤（绿网）'],
        ['文档生成', 'PPTX/DOCX 自动生成'],
    ],
    col_widths=[4, 10]
)

# ── 保存 ──
output_path = os.path.expanduser("~/projects/kg-edu/智课云枢_项目建设技术参数说明书_v2.0.docx")
doc.save(output_path)
print(f"✅ 文档已生成: {output_path}")
