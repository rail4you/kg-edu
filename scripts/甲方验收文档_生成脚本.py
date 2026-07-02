"""
课堂星智慧教学系统（第二期）—— 甲方验收签字文档
- 表格形式罗列所有需验收的功能
- 最后一列供甲方签字/盖章
- 不含内部技术细节，只列甲方可验证的功能要点
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell, color='666666', size='6'):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_run_font(run, name='Microsoft YaHei', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), name)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)
    if color:
        run.font.color.rgb = color


def add_para(doc, text, size=10.5, bold=False, color=None, align=None, space_after=2, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if align:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_section_title(doc, text, level=1):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(text)
    if level == 1:
        set_run_font(run, size=14, bold=True, color=RGBColor(0x16, 0x36, 0x5D))
    elif level == 2:
        set_run_font(run, size=12, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A))
    else:
        set_run_font(run, size=11, bold=True, color=RGBColor(0x42, 0x42, 0x42))


def add_data_row(table, cells_data, header=False, col_widths=None):
    row = table.add_row()
    cells = row.cells
    for i, data in enumerate(cells_data):
        cells[i].text = ''
        p = cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
        # 主内容
        r = p.add_run(data)
        if header:
            set_run_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            set_cell_bg(cells[i], '16365D')
        else:
            # 第一列加粗蓝色
            if i == 0:
                set_run_font(r, size=10, bold=True, color=RGBColor(0x16, 0x36, 0x5D))
            else:
                set_run_font(r, size=9.5)
        cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cells[i], color='888888', size='4')
        if col_widths and i < len(col_widths):
            cells[i].width = col_widths[i]


def build():
    doc = Document()

    # 全局字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rfonts.set(qn('w:ascii'), 'Microsoft YaHei')
    rfonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
    rpr.insert(0, rfonts)

    # 页边距
    section = doc.sections[0]
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # 横向页面（A4 横向更适合宽表）
    from docx.enum.section import WD_ORIENT
    new_w, new_h = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_w
    section.page_height = new_h

    # ============== 标题 ==============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('课堂星智慧教学系统（第二期）')
    set_run_font(r, size=22, bold=True, color=RGBColor(0x16, 0x36, 0x5D))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('功能验收清单（甲方签字版）')
    set_run_font(r, size=18, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A))

    # 元信息表
    info = doc.add_table(rows=2, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ['项目名称', '课堂星智慧教学系统（第二期）', '交付方', '研发交付组'],
        ['验收日期', '　　　年　　月　　日', '验收方', '　　　　　　　（盖章）'],
    ]
    for i, row_data in enumerate(info_data):
        cells = info.rows[i].cells
        for j, val in enumerate(row_data):
            cells[j].text = ''
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            if j % 2 == 0:
                set_run_font(r, size=10, bold=True, color=RGBColor(0x16, 0x36, 0x5D))
                set_cell_bg(cells[j], 'F0F4FA')
            else:
                set_run_font(r, size=10)
            set_cell_border(cells[j], color='16365D', size='4')
            cells[j].width = Cm(4)
    doc.add_paragraph()

    # ============== 说明 ==============
    add_section_title(doc, '验收说明', level=2)
    add_para(doc,
        '一、本验收清单严格基于《课堂星智慧教学系统功能参数（第二期）》八大功能模块制定，'
        '共 35 项功能条目，供甲方逐项确认。', size=10, indent=0.5)
    add_para(doc,
        '二、验收方式：建议甲方在测试环境（或本次交付环境）登录教师端、管理端、学生端，'
        '对照"验收要点"逐条操作验证。', size=10, indent=0.5)
    add_para(doc,
        '三、签字规则：每项功能"验收结论"列由甲方勾选 ✅通过 / ❌不通过 / ⚠有条件通过，'
        '并在最末"签字/盖章"列签名确认。', size=10, indent=0.5)
    add_para(doc,
        '四、若验收过程中发现问题，请于"备注"列简要描述（功能缺失 / 操作异常 / 体验不佳等），'
        '交付方将于约定时间内整改。', size=10, indent=0.5, space_after=6)

    # ====================================================================
    # 主验收表
    # ====================================================================
    # 列宽：序号 | 功能项 | 验收要点 | 验收方式 | 结论 | 签字/备注
    col_widths = [Cm(1.2), Cm(5.5), Cm(10.5), Cm(3.0), Cm(2.5), Cm(4.0)]

    main_table = doc.add_table(rows=1, cols=6)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    header_row = main_table.rows[0].cells
    for i, h in enumerate(['序号', '功能项', '验收要点', '验收方式', '结论', '签字/备注']):
        header_row[i].text = ''
        p = header_row[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(header_row[i], '16365D')
        header_row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(header_row[i], color='16365D', size='6')
        header_row[i].width = col_widths[i]

    # ===== 验收条目数据 =====
    items = []

    # -------- 模块一：知识图谱功能 --------
    items.append({'section': '一、知识图谱功能'})

    items.append({'idx': '1.1', 'name': '能力节点构建',
        'points': '① 手动创建能力节点 ② 批量导入能力节点 ③ AI 智能提取能力节点 ④ 节点含名称/描述/等级（初级/中级/高级） ⑤ 节点权重可设置',
        'method': '教师端 → 专业管理 → 能力图谱'})
    items.append({'idx': '1.2', 'name': '能力关系配置',
        'points': '① 支持能力间多对多关系 ② 关系类型：前置 / 并列 / 递进 ③ 关系线可标注说明文字',
        'method': '教师端 → 能力图谱 → 关系配置'})
    items.append({'idx': '1.3', 'name': '能力图谱展示',
        'points': '① 可视化展示能力节点与关系 ② 支持缩放、拖拽 ③ 节点/关系点击可查看详情',
        'method': '教师端 / 学生端'})
    items.append({'idx': '1.4', 'name': '思政知识点构建',
        'points': '① 手动创建思政知识点 ② 批量导入 ③ 节点含名称/描述/分类（思想道德修养/法律基础/党史教育等） ④ 节点权重可设置',
        'method': '教师端 → 思政图谱 → 知识点'})
    items.append({'idx': '1.5', 'name': '思政关系配置',
        'points': '① 支持思政知识点间关系配置 ② 关系线可标注说明',
        'method': '教师端 → 思政图谱 → 关系'})
    items.append({'idx': '1.6', 'name': '思政图谱展示',
        'points': '① 可视化展示思政知识点 ② 关联思政案例 ③ 案例详情可查看',
        'method': '教师端 / 学生端'})

    # -------- 模块二：AI 智能体功能 --------
    items.append({'section': '二、AI 智能体功能'})

    items.append({'idx': '2.1', 'name': 'AI 智能体生成（AI 出题）',
        'points': '① 针对指定知识点智能生成题目（选择/填空/简答等） ② 可自定义难度（1-5 级）③ 可自定义题型与数量 ④ 可参考已有内容出题 ⑤ 难度分级标注',
        'method': '教师端 → 练习 / 考试 → AI 出题'})
    items.append({'idx': '2.2', 'name': '教学素材生成',
        'points': '① AI 自动生成教案 ② AI 自动生成试题 ③ 素材可编辑、可下载',
        'method': '教师端 → AI 教学资料'})
    items.append({'idx': '2.3', 'name': 'AI 问答（基于知识图谱）',
        'points': '① 基于课程知识图谱的智能问答 ② 支持多轮对话与上下文理解 ③ 教师端可配置 Prompt ④ 学生端可获取即时解答',
        'method': '教师端 / 学生端 → AI 问答'})
    items.append({'idx': '2.4', 'name': 'AI 助教工具',
        'points': '① 一键生成 PPT ② 一键生成教学资料 ③ 辅助教学策略优化',
        'method': '教师端 → AI 助教 / 智慧空间'})
    items.append({'idx': '2.5', 'name': '智慧空间',
        'points': '① AI 智能体入口 ② AI 指令配置 ③ AI 工具集成 ④ 统一工作台',
        'method': '教师端 / 学生端 → 智慧空间'})

    # -------- 模块三：基础教学模块 --------
    items.append({'section': '三、基础教学模块'})

    items.append({'idx': '3.1', 'name': '课堂互动（在线讨论区）',
        'points': '① 实时在线讨论区 ② 师生互动 ③ 讨论记录可追溯',
        'method': '教师端 / 学生端 → 讨论区'})
    items.append({'idx': '3.2', 'name': '实时签到',
        'points': '① 教师发起签到 ② 学生端扫码或输入口令签到 ③ 出勤情况自动记录与统计',
        'method': '教师端 → 签到管理'})
    items.append({'idx': '3.3', 'name': '课堂测验（随堂测试）',
        'points': '① 教师可随时发起随堂测验 ② 学生即时答题 ③ 教师即时查看学生答题情况',
        'method': '教师端 → 课中测验'})
    items.append({'idx': '3.4', 'name': '考试管理',
        'points': '① 考试创建 / 发布 / 提交 ② 手动批改 + AI 智能批改 ③ 支持主观题（简答/论述）+ 客观题（选择） ④ AI 按知识点重难点出简单版/难题版试卷 ⑤ 随机生成试卷 ⑥ 成绩多维分析（成绩分布/知识点掌握率）',
        'method': '教师端 → 考试管理'})
    items.append({'idx': '3.5', 'name': '学习路径规划与进度跟踪',
        'points': '① 支持学习路径规划 ② 学生可查看学习进度 ③ 完成情况可追踪',
        'method': '学生端 → 我的学习'})
    items.append({'idx': '3.6', 'name': '混合式教学',
        'points': '① 支持交互课堂（教师与学生实时讨论） ② 实时签到功能 ③ 线上线下混合',
        'method': '教师端 → 课堂'})

    # -------- 模块四：学情分析模块 --------
    items.append({'section': '四、学情分析模块'})

    items.append({'idx': '4.1', 'name': '学生学习数据采集',
        'points': '① 采集学习行为数据 ② 采集考试成绩数据 ③ 支持多维度数据查询',
        'method': '管理端 → 系统日志'})
    items.append({'idx': '4.2', 'name': '班级整体学情分析',
        'points': '① 班级整体学情概览 ② 多维数据可视化（柱状图 / 饼图 / 折线图）③ 辅助教学决策',
        'method': '教师端 → 学情分析'})
    items.append({'idx': '4.3', 'name': '个性化学习建议',
        'points': '① 基于学习数据生成个性化建议 ② 推荐相关学习资源 ③ 优化学习路径',
        'method': '学生端 → 学习推荐'})

    # -------- 模块五：教学资源管理模块 --------
    items.append({'section': '五、教学资源管理模块'})

    items.append({'idx': '5.1', 'name': '资源推荐',
        'points': '① 基于学习行为智能推送资源 ② 基于知识点掌握情况推荐 ③ 资源分类清晰',
        'method': '学生端 → 推荐资源'})
    items.append({'idx': '5.2', 'name': '资源统计与分析',
        'points': '① 资源下载量统计 ② 观看次数统计 ③ 使用频率统计 ④ 资源热度可视化',
        'method': '教师端 / 管理端 → 资源统计'})

    # -------- 模块六：管理功能 --------
    items.append({'section': '六、管理功能'})

    items.append({'idx': '6.1', 'name': '学生管理',
        'points': '① 学生信息维护 ② 学生分组管理（按班级/小组） ③ 批量导入学生（含班级字段）',
        'method': '管理端 → 学生管理'})
    items.append({'idx': '6.2', 'name': '教师管理',
        'points': '① 教师信息维护 ② 含姓名 / 工号 / 职称 / 所属部门等基本信息',
        'method': '管理端 → 教师管理'})
    items.append({'idx': '6.3', 'name': '租户管理',
        'points': '① 租户级 AI 资源管理（AI Key 配置）② 租户内资源访问监控 ③ 系统状态总览',
        'method': '管理端 → 租户配置 / 系统监控'})
    items.append({'idx': '6.4', 'name': '密码管理（三级权限）',
        'points': '① 普通用户可自行修改密码 ② 学校管理员可重置师生密码 ③ 超级管理员可重置学校管理员密码 ④ 密码重置支持邮件通知',
        'method': '管理端 / 个人中心 → 密码管理'})
    items.append({'idx': '6.5', 'name': '用户分级与组织管理',
        'points': '① 支持学校 / 学院 / 班级多级组织 ② 组织内用户归属清晰',
        'method': '管理端 → 组织架构'})

    # -------- 模块七：附加功能 --------
    items.append({'section': '七、附加功能'})

    items.append({'idx': '7.1', 'name': '学生问题 → 教师邮箱',
        'points': '① 学生学习课程时可一键将问题发送至教师邮箱 ② 邮件标题/正文包含课程信息与问题内容 ③ 教师可邮件回复',
        'method': '学生端 → 课程 → 邮件问答'})
    items.append({'idx': '7.2', 'name': '多级知识点兼容（超星模板）',
        'points': '① 支持多级知识点结构 ② 兼容超星导出模板（OPML）③ 支持 Excel / XMind 多格式导入',
        'method': '教师端 → 知识点 → 导入'})
    items.append({'idx': '7.3', 'name': '自主可控：AI 服务层',
        'points': '① AI 服务层采用通义千问（阿里云）② 支持私有化部署，数据不出境 ③ 文件存储采用阿里云 OSS ④ AI Key 动态更换',
        'method': '管理端 → 系统配置'})
    items.append({'idx': '7.4', 'name': '自主可控：国产化技术栈',
        'points': '① 前端：Ant Design（蚂蚁集团）+ ECharts（百度）② 构建：Vite（尤雨溪）③ AI Agent / API：基于通义千问 API ④ 支持国产芯片（龙芯/鲲鹏/飞腾）适配',
        'method': '管理端 → 系统信息'})
    items.append({'idx': '7.5', 'name': '数据迁移',
        'points': '① 课程数据可按相关接口标准导出 ② 互联网平台 → 专网平台迁移便捷 ③ 支持租户级数据备份与恢复',
        'method': '管理端 → 数据备份'})

    # -------- 模块八：微专业 --------
    items.append({'section': '八、微专业'})

    items.append({'idx': '8.1', 'name': '微专业管理',
        'points': '① 微专业创建与发布 ② 课程内容编排（视频/资源/章节）③ 学生选课 ④ 作业 / 练习 / 学习进度跟踪 ⑤ 教师批改',
        'method': '教师端 / 学生端 → 微专业'})

    # 写入表格
    for item in items:
        if 'section' in item:
            # 模块标题行（合并 6 列）
            row = main_table.add_row()
            cells = row.cells
            merged = cells[0].merge(cells[5])
            merged.text = ''
            p = merged.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.3)
            r = p.add_run(item['section'])
            set_run_font(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            set_cell_bg(merged, '2E5C8A')
            set_cell_border(merged, color='16365D', size='4')
            merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            row = main_table.add_row().cells
            data = [
                item['idx'],
                item['name'],
                item['points'],
                item['method'],
                '☐ 通过\n☐ 不通过\n☐ 有条件',
                ''
            ]
            for i, val in enumerate(data):
                row[i].text = ''
                p = row[i].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                if i in [0, 3, 4]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val)
                if i == 0:
                    set_run_font(r, size=10, bold=True, color=RGBColor(0x16, 0x36, 0x5D))
                elif i == 1:
                    set_run_font(r, size=10, bold=True)
                else:
                    set_run_font(r, size=9.5)
                row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_border(row[i], color='888888', size='4')
                row[i].width = col_widths[i]

    # ============== 验收结论汇总 ==============
    doc.add_paragraph()
    add_section_title(doc, '验收结论与签字', level=1)

    # 验收结论表
    result_table = doc.add_table(rows=5, cols=2)
    result_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    result_data = [
        ('总体验收结论', '☐ 全部通过，验收合格\n☐ 部分通过，需整改后复验\n☐ 未通过，需重新交付'),
        ('通过项数量', '　　　项 / 共 35 项'),
        ('不通过项数量', '　　　项'),
        ('整改要求（概要）', ''),
        ('复验日期', '　　　年　　月　　日'),
    ]
    for i, (label, value) in enumerate(result_data):
        cells = result_table.rows[i].cells
        cells[0].text = ''
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(label)
        set_run_font(r, size=11, bold=True, color=RGBColor(0x16, 0x36, 0x5D))
        set_cell_bg(cells[0], 'F0F4FA')
        set_cell_border(cells[0], color='16365D', size='4')
        cells[0].width = Cm(5)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        cells[1].text = ''
        p = cells[1].paragraphs[0]
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        set_run_font(r, size=11)
        set_cell_border(cells[1], color='16365D', size='4')
        cells[1].width = Cm(20)
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # 签字栏
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_data = [
        ('交付方（盖章）', '验收方（盖章）'),
        ('', ''),
    ]
    for i, (left, right) in enumerate(sig_data):
        cells = sig_table.rows[i].cells
        for j, val in enumerate([left, right]):
            cells[j].text = ''
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0:
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(val)
                set_run_font(r, size=12, bold=True, color=RGBColor(0x16, 0x36, 0x5D))
                set_cell_bg(cells[j], 'F0F4FA')
            else:
                # 留空签字区
                p.paragraph_format.space_after = Pt(60)  # 给签字留空
                r = p.add_run('\n\n\n\n\n')  # 多个换行留出签字空间
                set_run_font(r, size=10)
                p2 = cells[j].add_paragraph('签字人：　　　　　　　　　日期：　　　年　　月　　日')
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(p2.runs[0], size=10)
            set_cell_border(cells[j], color='16365D', size='6')
            cells[j].width = Cm(12.5)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 页脚说明
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('本验收单一式两份，甲乙双方各执一份，具有同等法律效力。')
    set_run_font(r, size=9, color=RGBColor(0x88, 0x88, 0x88))

    # 保存
    out_path = '/Users/bai/projects/kg-edu/docs/课堂星智慧教学系统第二期_甲方验收签字版.docx'
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f'✅ 甲方验收文档已生成: {out_path}')
    print(f'   文件大小: {os.path.getsize(out_path) / 1024:.1f} KB')
    return out_path


if __name__ == '__main__':
    build()
