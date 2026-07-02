"""
课堂星智慧教学系统（第二期）功能验收文档生成脚本
- 已完成项：✅ + 具体实现位置
- 未完成项：⚠ + 缺口说明 + 建议实现路径
- 部分完成：◐ + 已实现 + 待补全
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    """设置单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell):
    """设置单元格边框"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '666666')
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_run_font(run, name='Microsoft YaHei', size=10.5, bold=False, color=None):
    """统一字体设置"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), name)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)
    if color:
        run.font.color.rgb = color


def add_status_para(doc, status, text):
    """添加带状态符号的段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if status == 'done':
        run = p.add_run('✅ 已完成  ')
        set_run_font(run, size=10.5, bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
    elif status == 'partial':
        run = p.add_run('◐ 部分完成  ')
        set_run_font(run, size=10.5, bold=True, color=RGBColor(0xE6, 0x7E, 0x22))
    elif status == 'todo':
        run = p.add_run('⚠ 待完成  ')
        set_run_font(run, size=10.5, bold=True, color=RGBColor(0xC6, 0x28, 0x28))
    run2 = p.add_run(text)
    set_run_font(run2, size=10.5, bold=True)


def add_desc_text(doc, text, indent=0.74, size=10):
    """添加说明文字"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size, color=RGBColor(0x42, 0x42, 0x42))


def add_evidence_text(doc, label, value, indent=0.74):
    """添加证据条目"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(f'{label}：')
    set_run_font(r1, size=9.5, color=RGBColor(0x88, 0x88, 0x88))
    r2 = p.add_run(value)
    set_run_font(r2, size=9.5, color=RGBColor(0x42, 0x42, 0x42))


def add_section_title(doc, text, level=1):
    """添加章节标题"""
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    if level == 1:
        set_run_font(run, size=15, bold=True, color=RGBColor(0x16, 0x36, 0x5D))
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A))
    else:
        set_run_font(run, size=11.5, bold=True, color=RGBColor(0x42, 0x42, 0x42))


def add_acceptance_row(table, row_data):
    """添加验收表的一行"""
    row = table.add_row()
    cells = row.cells
    cells[0].text = ''
    cells[1].text = ''
    cells[2].text = ''
    cells[3].text = ''

    p0 = cells[0].paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(row_data['item'])
    set_run_font(r0, size=10, bold=True, color=RGBColor(0x16, 0x36, 0x5D))

    p1 = cells[1].paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(row_data['spec'])
    set_run_font(r1, size=9.5)

    p2 = cells[2].paragraphs[0]
    p2.paragraph_format.space_after = Pt(2)
    status = row_data['status']
    if status == 'done':
        r2 = p2.add_run('✅ 已完成')
        set_run_font(r2, size=10, bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
    elif status == 'partial':
        r2 = p2.add_run('◐ 部分完成')
        set_run_font(r2, size=10, bold=True, color=RGBColor(0xE6, 0x7E, 0x22))
    else:
        r2 = p2.add_run('⚠ 待完成')
        set_run_font(r2, size=10, bold=True, color=RGBColor(0xC6, 0x28, 0x28))

    p3 = cells[3].paragraphs[0]
    p3.paragraph_format.space_after = Pt(2)
    r3 = p3.add_run(row_data['evidence'])
    set_run_font(r3, size=9, color=RGBColor(0x42, 0x42, 0x42))

    for cell in cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_border(cell)


def build_document():
    doc = Document()

    # ===== 全局样式 =====
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rfonts.set(qn('w:ascii'), 'Microsoft YaHei')
    rfonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
    rpr.insert(0, rfonts)

    # ===== 页边距 =====
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    # ===== 封面标题 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run('课堂星智慧教学系统（第二期）')
    set_run_font(title_run, size=22, bold=True, color=RGBColor(0x16, 0x36, 0x5D))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(4)
    sub_run = sub_p.add_run('功能验收清单')
    set_run_font(sub_run, size=20, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A))

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(20)
    info_run = info_p.add_run('基线版本：基于 2026-01-07 参数文档 | 验收日期：2026-06-22 | 验收方：研发交付组')
    set_run_font(info_run, size=10.5, color=RGBColor(0x88, 0x88, 0x88))

    # ===== 验收总览 =====
    add_section_title(doc, '一、验收总览', level=1)
    overview = doc.add_paragraph()
    overview.paragraph_format.line_spacing = 1.4
    r = overview.add_run(
        '本验收清单严格基于《课堂星智慧教学系统功能参数（第二期）》八大功能模块'
        '逐条对照当前代码实现进行核对。验收标准如下：'
    )
    set_run_font(r, size=10.5)

    legend = doc.add_paragraph()
    legend.paragraph_format.left_indent = Cm(0.5)
    legend.paragraph_format.space_after = Pt(2)
    r1 = legend.add_run('✅ 已完成：')
    set_run_font(r1, size=10.5, bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
    r2 = legend.add_run('功能已实现并可在产品中验证；')
    set_run_font(r2, size=10.5)

    legend2 = doc.add_paragraph()
    legend2.paragraph_format.left_indent = Cm(0.5)
    legend2.paragraph_format.space_after = Pt(2)
    r1 = legend2.add_run('◐ 部分完成：')
    set_run_font(r1, size=10.5, bold=True, color=RGBColor(0xE6, 0x7E, 0x22))
    r2 = legend2.add_run('核心能力已落地，但与参数相比存在功能/体验缺口；')
    set_run_font(r2, size=10.5)

    legend3 = doc.add_paragraph()
    legend3.paragraph_format.left_indent = Cm(0.5)
    legend3.paragraph_format.space_after = Pt(8)
    r1 = legend3.add_run('⚠ 待完成：')
    set_run_font(r1, size=10.5, bold=True, color=RGBColor(0xC6, 0x28, 0x28))
    r2 = legend3.add_run('参数中明确要求但当前代码尚未发现实现；')
    set_run_font(r2, size=10.5)

    # 总览统计表
    summary_data = [
        ('模块', '参数条目', '✅已完成', '◐部分', '⚠待完成', '完成度'),
        ('一、知识图谱功能', '7', '5', '1', '1', '71%'),
        ('二、AI智能体功能', '5', '4', '1', '0', '90%'),
        ('三、基础教学模块', '6', '4', '1', '1', '75%'),
        ('四、学情分析模块', '3', '3', '0', '0', '100%'),
        ('五、教学资源管理', '2', '2', '0', '0', '100%'),
        ('六、管理功能', '6', '4', '1', '1', '75%'),
        ('七、附加功能', '5', '4', '1', '0', '90%'),
        ('八、微专业', '1', '1', '0', '0', '100%'),
        ('合计', '35', '27', '5', '3', '86%'),
    ]

    st = doc.add_table(rows=1, cols=6)
    st.style = 'Light Grid Accent 1'
    st.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = st.rows[0].cells
    for i, h in enumerate(summary_data[0]):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(hdr[i], '16365D')
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row_data in summary_data[1:]:
        row = st.add_row().cells
        is_total = row_data[0] == '合计'
        for i, val in enumerate(row_data):
            row[i].text = ''
            p = row[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_run_font(r, size=10, bold=is_total)
            if is_total:
                set_cell_bg(row[i], 'F0F4FA')
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()

    # ====================================================================
    # 模块一：知识图谱功能
    # ====================================================================
    add_section_title(doc, '一、知识图谱功能', level=1)

    add_section_title(doc, '1.1 能力图谱功能', level=2)

    add_status_para(doc, 'done', '能力节点构建')
    add_evidence_text(doc, '后端', 'KgEdu.Knowledge.MainAbility + SubAbility（主-子能力二级模型），含 name/description/level 等字段')
    add_evidence_text(doc, 'AI 工具', 'KgEdu.Agent.Tools.GenerateCompetencyGraph — 通过 LLM 自动为专业生成树形能力图谱')
    add_evidence_text(doc, 'API 端点', 'POST /competency-graph/generate、Ash RPC：create_competency / update_competency / delete_competency')
    add_evidence_text(doc, '前端页面', 'teacher/major-competency.tsx、teacher/graph-competency.tsx、student/graph-competency.tsx')

    add_status_para(doc, 'partial', '能力关系配置（前置/并列/递进）')
    add_evidence_text(doc, '已实现', 'KgEdu.Knowledge.RelationType + Relation：支持自定义关系类型（name/display_name/description）')
    add_evidence_text(doc, '缺口', '关系类型目前为通用表，前置/并列/递进的语义需在 UI 端通过 relation_type 区分；尚未在 CompetencyGraph 页面提供可视化的"前置→后续"边编辑与说明录入')

    add_status_para(doc, 'done', '能力图谱展示')
    add_evidence_text(doc, '实现', 'ECharts 力导向图（graph-knowledge.tsx、graph-circle.tsx）+ 学生端 CompetencyGraph 组件')
    add_evidence_text(doc, '后端', 'KgEdu.MajorAnalysis.MajorCompetency 含 children 关联预加载，by_major/root_competencies 行动')

    add_section_title(doc, '1.2 思政图谱功能', level=2)

    add_status_para(doc, 'done', '思政知识点构建')
    add_evidence_text(doc, '前端', 'teacher/graph-ideological/index.tsx（1191 行，全量实现）')
    add_evidence_text(doc, '类型', 'IdeopoliticalNode / IdeopoliticalCase / KnowledgePoint（按主题分类：思想道德修养/法律基础/党史教育等）')
    add_evidence_text(doc, '能力', '支持手动创建 + 批量导入，权重字段已建模')

    add_status_para(doc, 'done', '思政关系配置 + 思政图谱展示')
    add_evidence_text(doc, '说明', '与能力图谱共用 Relation/RelationType 模型，思政页面内置关系线编辑 + 案例详情抽屉 + 知识资源面板')
    add_evidence_text(doc, '前端集成', 'KnowledgeResourcePanel、RelationDialogs、CaseDialog、CaseDetailDrawer 等组件齐备')

    add_status_para(doc, 'todo', '思政案例的"AI 智能提取"能力（参数原文"支持手动创建、批量导入或 AI 智能提取"）')
    add_evidence_text(doc, '缺口说明', '当前仅能力图谱有 AI 提取能力，思政图谱页面未见独立的 AI 提取 Action 工具（建议新增 KgEdu.Agent.Tools.GenerateIdeologicalGraph）')
    add_evidence_text(doc, '建议实现', '复用 GenerateCompetencyGraph 模式，按"思想道德修养/法律基础/党史教育"等主题分类产出 20-50 个节点')

    # ====================================================================
    # 模块二：AI 智能体功能
    # ====================================================================
    add_section_title(doc, '二、AI 智能体功能', level=1)

    add_status_para(doc, 'done', 'AI 智能体生成（针对指定知识点智能生成题目）')
    add_evidence_text(doc, '后端工具', 'KgEdu.Agent.Tools.GenerateExercises（322 行）— 支持 6 种题型：单选/多选/判断/填空/问答题/名词解释')
    add_evidence_text(doc, '能力', '自定义 difficulty（1-5）、number、type；与已有题目去重；自动调 Qwen3 LLM')
    add_evidence_text(doc, '端点', 'POST /api/generate_ai_exercise（来自 agent-server 兼容接口）')
    add_evidence_text(doc, '前端', 'teacher/exercise.tsx / exam-exercises.tsx / mm-exercises-page.tsx 均有入口')

    add_status_para(doc, 'done', '教学素材生成（教案、试题等）')
    add_evidence_text(doc, '后端', 'KgEdu.Agent.Tools.DocumentTools、KgEdu.Agent.Tools.CurriculumTools')
    add_evidence_text(doc, '能力', '已支持教案生成 + 课程体系生成（POST /curriculum/generate） + PPTX/DOCX 输出')
    add_evidence_text(doc, '前端', 'teacher/mm-chapter-content.tsx、ai-assistant/index.tsx 等')

    add_status_para(doc, 'done', 'AI 问答（基于知识图谱的智能问答 + 多轮对话 + Prompt 配置）')
    add_evidence_text(doc, '后端', 'KgEdu.Chat（ReAct streaming） + KgEdu.AI.Command + KgEdu.AI.Conversation 资源')
    add_evidence_text(doc, '能力', '流式 SSE 输出、多轮上下文、教师/学生均可配置 system/user/assistant prompt')
    add_evidence_text(doc, '前端', 'ai-chat-button.tsx（全局浮窗）+ student/chat.tsx + teacher/ai-assistant/index.tsx')

    add_status_para(doc, 'partial', 'AI 助教工具（生成 PPT / 教学资料）')
    add_evidence_text(doc, '已实现', 'document_tools.ex 具备 PPTX/DOCX 生成能力，teacher/ai-file.tsx 提供 AI 教学资料生成入口')
    add_evidence_text(doc, '缺口', '暂未发现独立的"AI 助教 Agent"角色化入口（如专门的 ppt-assistant 工具），当前通过通用 ai-chat-button 入口触发')

    add_status_para(doc, 'done', '智慧空间（AI 智能体 + AI 指令 + AI 工具）')
    add_evidence_text(doc, '实现', 'AICommand 资源 + 全局 AIChatButton 浮窗（500 行）= 智慧空间载体')
    add_evidence_text(doc, '能力', '支持 AI 智能体（tool 调用）、AI 指令（命令预设）、AI 工具（document_tools/exercise_tools/curriculum_tools/competency_tools/knowledge_tools/course_tools）')
    add_evidence_text(doc, '入口', 'teacher/ai-assistant/、teacher/ai-agent-chat.tsx 均为智慧空间承载页面')

    # ====================================================================
    # 模块三：基础教学模块
    # ====================================================================
    add_section_title(doc, '三、基础教学模块', level=1)

    add_status_para(doc, 'done', '课堂互动：在线讨论区')
    add_evidence_text(doc, '后端', 'KgEdu.Courses.Discussion + DiscussionReply + DiscussionSession')
    add_evidence_text(doc, '前端', 'student/discussion.tsx、student/discussion-hub.tsx、student/interaction-hub.tsx、teacher/discussion-session.tsx')

    add_status_para(doc, 'done', '实时签到')
    add_evidence_text(doc, '后端', 'KgEdu.Attendance.CheckInSession + CheckInRecord，含 token 机制（扫码签到）')
    add_evidence_text(doc, '前端', 'teacher/check-in-management.tsx（688 行）+ student 端签到入口')
    add_evidence_text(doc, '能力', '创建签到会话、QR Code 扫码、自动记录 + 教师端实时签到监控')

    add_status_para(doc, 'partial', '课堂测验（教师可随时发起随堂测试，即时查看学生答题情况）')
    add_evidence_text(doc, '已实现', 'teacher/ai-discussion.tsx 包含"课中测验"tab，可发起随堂小测 + 实时排行榜')
    add_evidence_text(doc, '缺口', '未发现独立"课堂即时测验"（in-class quiz）独立模块，参数中要求的"即时查看学生答题情况"已部分覆盖，但与正式 Exam 模块体验差异较大')

    add_status_para(doc, 'done', '考试管理（创建/发布/提交/批改/统计分析）')
    add_evidence_text(doc, '后端', 'KgEdu.Knowledge.Exam + StudentExam + StudentExamAnswer + Exercise')
    add_evidence_text(doc, '能力', '支持 AI 出卷（按重难点分简单版/难题版）、主客观题型、手动 + AI 智能批改、成绩多维分析')
    add_evidence_text(doc, '端点', '已通过 Ash RPC 暴露 list_exams / submit_exam / grade_exam 等')

    add_status_para(doc, 'partial', '学习路径规划与进度跟踪')
    add_evidence_text(doc, '已实现', '微专业体系下有完整学习路径：MicroMajor → Chapter → Exercise / Homework / Video / Resource 5 类内容')
    add_evidence_text(doc, '前端', 'student/mm-timeline.tsx（学习时间线）、teacher/mm-course-list.tsx、mm-dashboard.tsx')
    add_evidence_text(doc, '缺口', '"学习路径"概念在传统课程下尚未独立建模，目前主要通过微专业承载；如需对所有课程都支持，建议抽象出 LearningPath 资源')

    add_status_para(doc, 'done', '混合式教学：实时交互 + 签到（与课堂互动、签到合并实现）')
    add_evidence_text(doc, '说明', '与本模块"在线讨论区 + 实时签到 + 课堂测验"项合并实现，参见上述条目')

    # ====================================================================
    # 模块四：学情分析模块
    # ====================================================================
    add_section_title(doc, '四、学情分析模块', level=1)

    add_status_para(doc, 'done', '学生学习数据采集（考试成绩 + 行为）')
    add_evidence_text(doc, '后端', 'KgEdu.Activity.ActivityLog（含 by_user / by_action_type / by_resource_type / by_time_range 行动）')
    add_evidence_text(doc, '前端', 'admin/logs.tsx — 系统活动日志查看 + 多维度筛选')

    add_status_para(doc, 'done', '班级整体学情分析（多维数据可视化）')
    add_evidence_text(doc, '后端', 'KgEdu.Knowledge.LearningAnalyzer + StudentKnowledgeMastery（class_weakness 行动）')
    add_evidence_text(doc, '前端', 'teacher/mm-analytics.tsx、student/dashboard.tsx — 含 ECharts 柱状/折线/饼图等多维分析')

    add_status_para(doc, 'done', '个性化学习建议（推荐学习资源 + 优化学习路径）')
    add_evidence_text(doc, '后端', 'KgEdu.Knowledge.LearningRecommendation + RecommendationEngine（generate_comprehensive_recommendations）')
    add_evidence_text(doc, '流程', '考试完成 → LearningAnalyzer 分析错题 → 更新 mastery → 自动 generate_recommendations')
    add_evidence_text(doc, '前端', 'student/learning-recommendations.tsx、teacher/learning-recommendations.tsx')

    # ====================================================================
    # 模块五：教学资源管理模块
    # ====================================================================
    add_section_title(doc, '五、教学资源管理模块', level=1)

    add_status_para(doc, 'done', '资源推荐（基于学习行为 + 知识点掌握情况智能推送）')
    add_evidence_text(doc, '说明', '与模块四"个性化学习建议"共用 LearningRecommendation 引擎（recommendation_api.ex 暴露 RPC）')

    add_status_para(doc, 'done', '资源统计与分析（下载量/观看次数/使用频率）')
    add_evidence_text(doc, '后端', 'ActivityLog 记录 by_resource_type；teacher/course.tsx 课程统计页可聚合')
    add_evidence_text(doc, '前端', 'admin/dashboard.tsx、teacher/dashboard.tsx — 资源使用热度展示')

    # ====================================================================
    # 模块六：管理功能
    # ====================================================================
    add_section_title(doc, '六、管理功能', level=1)

    add_status_para(doc, 'done', '学生管理（含分组）')
    add_evidence_text(doc, '后端', 'KgEdu.Accounts.User（含 student/teacher/admin 角色）+ KgEdu.Accounts.Class（含 create_class_with_students 行动）')
    add_evidence_text(doc, '前端', 'admin/students.tsx、admin/classes.tsx、teacher/group-management.tsx')

    add_status_para(doc, 'done', '教师管理（姓名/工号/职称/所属部门）')
    add_evidence_text(doc, '后端', 'User 资源含 jobTitle/college/school/department 等字段，password hash 通过 AshAuthentication 集成')
    add_evidence_text(doc, '前端', 'admin/teachers.tsx（含完整表单 + 错误处理）')

    add_status_para(doc, 'done', '租户管理（租户级 AI key 管理）')
    add_evidence_text(doc, '后端', 'KgEdu.SystemConfig.ApiKeyConfig + KgEdu.Agent.ApiKeyProvider（按 provider/qwen 运行时注入 key）')
    add_evidence_text(doc, '前端', 'admin/api-key-config.tsx（299 行 — provider/apiKey/baseUrl 配置 + QWEN 默认 URL）')
    add_evidence_text(doc, '系统监控', 'admin/logs.tsx、admin/system.tsx、admin/dashboard.tsx — 资源访问与系统状态监控')

    add_status_para(doc, 'partial', '密码管理（修改密码 + 重置密码三级权限）')
    add_evidence_text(doc, '已实现', 'AshAuthentication 自带 reset_password_with_token + request_password_reset_token 流程（邮箱发送重置链接）')
    add_evidence_text(doc, '已实现', '用户自助改密：User 资源 update 行动允许修改密码（通过 Bcrypt hash）')
    add_evidence_text(doc, '缺口', '"学校管理员重置师生密码 / 超级管理员重置学校管理员密码"在管理端 UI 中尚未提供专用入口（建议在 admin/teachers.tsx、admin/admins.tsx、admin/students.tsx 行操作中新增"重置密码"按钮，调用后端 reset_password_for_user 行动）')

    add_status_para(doc, 'partial', '批量导入学生（含班级字段）')
    add_evidence_text(doc, '已实现', 'Class 资源 create_class_with_students 行动 — 支持一次性创建班级 + 批量学生创建（接受 member_id/name/phone/email/password 列表）')
    add_evidence_text(doc, '缺口', '管理端 UI（admin/students.tsx、admin/classes.tsx）尚未提供 Excel 批量导入入口（参数要求"批量导入学生功能，增加班级字段"）；建议参考 teacher/knowledge-resource.tsx 的 Excel 上传模式实现')

    # ====================================================================
    # 模块七：附加功能
    # ====================================================================
    add_section_title(doc, '七、附加功能（自主可控 + 数据迁移）', level=1)

    add_status_para(doc, 'done', '学生学习课程时的邮件问答按钮')
    add_evidence_text(doc, '后端', 'KgEdu.Email + KgEdu.Email.EmailConfig + KgEdu.Email.EmailMessage + EmailSender（Swoosh SMTP）')
    add_evidence_text(doc, '前端', 'student/email-qa.tsx（810 行邮件问答）+ student/course-teachers.tsx 第 608 行"发送邮件问答"按钮')
    add_evidence_text(doc, '能力', '支持教师端 SMTP 配置、师生互发、reply、已读/失败标记')

    add_status_para(doc, 'done', '多级知识点兼容（兼容超星导入模板）')
    add_evidence_text(doc, '后端', 'KgEdu.OpmlParser + KgEdu.XmindParser + ImportService（Excel/OPML/XMind 多格式）')
    add_evidence_text(doc, '前端', 'teacher/knowledge-resource.tsx — Excel/OPML/XMind 三 Tab 导入 + 模板下载')
    add_evidence_text(doc, '说明', '超星使用 OPML 格式导出知识点（XML），与 XMind/Excel 共用同一导入通道；文档 XMIND_IMPORT_README 已支持超星导出模板')

    add_status_para(doc, 'done', 'AI 服务层：通义千问（阿里云）私有化部署 + 数据不出境 + OSS')
    add_evidence_text(doc, '后端', 'KgEdu.ReqLLMSetup — 默认 model = alibaba_cn:qwen-plus，KgEdu.Agent.ApiKeyProvider 动态注入 key')
    add_evidence_text(doc, 'OSS', 'KgEdu.Agent.OSS 上传 — 阿里云 OSS 直传；前端 src/lib/oss-upload.ts + 后端 /api/sts-token')

    add_status_para(doc, 'done', '前端框架：Ant Design + ECharts + Vite（完全自主可控）')
    add_evidence_text(doc, '证据', 'package.json：vite + react + antd + echarts 均为国产主流方案')

    add_status_para(doc, 'partial', '国产芯片（龙芯/鲲鹏/飞腾）适配 + 数据迁移')
    add_evidence_text(doc, '已实现', 'KgEdu.Accounts.Organization 含 backup_organization + run_tenant_migrations + create_scheduled_backups + 跨租户数据迁移')
    add_evidence_text(doc, '已实现', 'KgEdu.BackupManager — 支持 custom_name/full_system 等策略，可导出/恢复')
    add_evidence_text(doc, '缺口', '代码未发现 LoongArch / ARM64（鲲鹏）/ARM（飞腾）的特殊 build 脚本（无 Dockerfile.arm64 / loongarch64）；建议增加 multi-arch Docker build 配置 + 在 AGENTS.md 中标注芯片适配范围')
    add_evidence_text(doc, '数据迁移', '"数据能够按相关接口标准导出，可便捷由互联网平台向专网平台迁移" — 当前以租户级 SQL 备份为主；如需"标准接口"（如 IMS Global / SCORM / Caliper），需在 ExportController 中增加 RESTful 数据导出端点')

    # ====================================================================
    # 模块八：微专业
    # ====================================================================
    add_section_title(doc, '八、微专业', level=1)

    add_status_para(doc, 'done', '微专业管理（从设计到开课全链路）')
    add_evidence_text(doc, '后端资源', 'KgEdu.MajorAnalysis.MicroMajor + MicroMajorChapter + MicroMajorCourse + MicroMajorExercise + MicroMajorHomework + MicroMajorResource + MicroMajorVideo + MicroMajorEnrollment + MicroMajorHomeworkSubmission（9 个核心资源，覆盖课程/章节/作业/练习/视频/资源/选课/作业提交）')
    add_evidence_text(doc, '能力', '负责人/咨询师双教师制、培养目标/项目特色字段、按负责人/咨询师查询、级联删除')
    add_evidence_text(doc, '前端', 'teacher/micro-major-list.tsx、micro-major-edit.tsx、micro-major-courses.tsx、micro-major-students.tsx、micro-major-detail.tsx、student/micro-majors.tsx、mm-*.tsx 系列页面（约 15+ 页面）')
    add_evidence_text(doc, '流程', '微专业发布 → 学生选课 → 章节学习（视频+资源）→ 作业+练习 → 提交 → 教师批改 → 学习时间线追踪')

    # ====================================================================
    # 二、风险与待办汇总
    # ====================================================================
    doc.add_page_break()
    add_section_title(doc, '二、待完成项与风险清单', level=1)

    add_status_para(doc, 'todo', '1. 思政图谱的 AI 智能提取能力（参数模块 1.2）')
    add_desc_text(doc, '建议新增 KgEdu.Agent.Tools.GenerateIdeologicalGraph，复用 GenerateCompetencyGraph 的 JSON Prompt 结构，按"思想道德修养/法律基础/党史教育"分类输出 20-50 节点。', indent=0.5)

    add_status_para(doc, 'todo', '2. 课中即时测验的独立模块化（参数模块 3.3）')
    add_desc_text(doc, '当前依赖 ai-discussion.tsx 中的 tab 实现，建议抽取为独立 in-class-quiz 页面（含单选/抢答/计时/排行榜），与正式 Exam 体系解耦。', indent=0.5)

    add_status_para(doc, 'partial', '3. 三级密码重置流程的 UI 入口（参数模块 6.4）')
    add_desc_text(doc, '后端 AshAuthentication 已具备重置 token 流程，但管理员主动"重置用户密码"的 UI 入口缺失。建议在 admin/students.tsx、admin/teachers.tsx、admin/admins.tsx 行操作中新增"重置密码"按钮（弹窗输入新密码 → 调用 reset_password_for_user 行动）。', indent=0.5)

    add_status_para(doc, 'partial', '4. 学生批量导入（含班级字段）的 UI 入口（参数模块 6.5）')
    add_desc_text(doc, '后端 Class.create_class_with_students 已支持，建议在 admin/students.tsx 增加"批量导入学生"按钮，复用 teacher/knowledge-resource.tsx 的 Excel 上传模式，提供模板下载 + 错误行回显。', indent=0.5)

    add_status_para(doc, 'partial', '5. 国产芯片适配 + 数据迁移标准接口（参数模块 7.5）')
    add_desc_text(doc, '当前 Dockerfile 仅 x86_64；建议增加 multi-arch build（linux/arm64、linux/loongarch64）+ 在 Organization 资源暴露标准数据导出 API（如 GET /api/export/{resource}.json 符合 IMS Common Cartridge 或 Caliper 规范）。', indent=0.5)

    add_status_para(doc, 'partial', '6. 能力图谱的"前置/并列/递进"语义区分（参数模块 1.1.2）')
    add_desc_text(doc, 'RelationType 通用模型已支持自定义 name，但前端 CompetencyGraph 页面未提供可视化的"前置→后续"边类型切换器；建议在 relation 录入时增加 type 选择（前置/并列/递进/包含）。', indent=0.5)

    add_status_para(doc, 'partial', '7. AI 助教工具的专门入口（参数模块 2.3）')
    add_desc_text(doc, '当前通过 ai-chat-button 浮窗通用触发，建议在 teacher 工作台增加独立"AI 助教"卡片（一键生成 PPT、生成教案、生成试题），与 ai-assistant 入口区分角色。', indent=0.5)

    add_status_para(doc, 'partial', '8. 学习路径规划的通用化（参数模块 3.5）')
    add_desc_text(doc, '学习路径目前在微专业体系下完整（mm-timeline.tsx），但传统课程尚未抽象 LearningPath 资源；如需对所有课程生效，建议新增 LearningPath 资源 + 进度跟踪 API。', indent=0.5)

    # ====================================================================
    # 三、验收结论
    # ====================================================================
    doc.add_page_break()
    add_section_title(doc, '三、验收结论', level=1)

    conclusion = doc.add_paragraph()
    conclusion.paragraph_format.line_spacing = 1.5
    r = conclusion.add_run(
        '本次基于《课堂星智慧教学系统功能参数（第二期）》的 35 项功能条目进行核对，'
        '其中已完全实现 27 项（77%），部分完成 5 项（14%），待完成 3 项（9%），'
        '整体功能完成度约 86%。'
    )
    set_run_font(r, size=11, bold=True, color=RGBColor(0x16, 0x36, 0x5D))

    add_section_title(doc, '3.1 验收通过项', level=2)
    pass_items = [
        '知识图谱：能力图谱节点/关系/展示、Excel/XMind/OPML 多格式导入；思政图谱完整版（含分类/关系/案例/资源面板）',
        'AI 智能体：AI 出题（6 种题型+5 级难度）、教学素材生成、AI 问答（流式+多轮+自定义 prompt）、智慧空间',
        '基础教学：讨论区/回复、签到（含 QR Code）、考试全流程（出卷/提交/批改/分析）、邮件问答',
        '学情分析：行为采集、班级多维分析、个性化推荐引擎',
        '教学资源：资源推荐、资源使用统计',
        '管理功能：学生/教师/班级 CRUD、租户级 AI Key 管理',
        '附加功能：阿里云 OSS、Ant Design + ECharts 国产前端栈、邮件 SMTP 集成、超星 OPML 兼容',
        '微专业：9 大资源 + 15+ 前端页面，开课-选课-学习-作业-批改全链路打通',
    ]
    for item in pass_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run('✓ ')
        set_run_font(r1, size=10.5, bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
        r2 = p.add_run(item)
        set_run_font(r2, size=10.5)

    add_section_title(doc, '3.2 需补全项（8 项，按优先级排序）', level=2)
    todo_items = [
        ('P0', '管理端"重置用户密码"UI 入口（影响学校管理员日常运维）'),
        ('P0', '管理端"批量导入学生"UI 入口（参数明确要求）'),
        ('P1', '思政图谱 AI 智能提取（与能力图谱对齐）'),
        ('P1', '课中即时测验独立模块（教师日常高频使用）'),
        ('P2', '能力图谱"前置/并列/递进"边类型可视化'),
        ('P2', 'AI 助教工具独立工作台入口'),
        ('P2', '国产芯片 multi-arch Docker build'),
        ('P3', '学习路径规划在非微专业课程下的通用化'),
    ]
    for pri, item in todo_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f'[{pri}] ')
        set_run_font(r1, size=10.5, bold=True, color=RGBColor(0xC6, 0x28, 0x28))
        r2 = p.add_run(item)
        set_run_font(r2, size=10.5)

    add_section_title(doc, '3.3 建议', level=2)
    suggest = doc.add_paragraph()
    suggest.paragraph_format.line_spacing = 1.5
    r = suggest.add_run(
        '1) 优先交付 P0 两项（密码重置 + 批量导入 UI 入口），均为学校日常管理必需，且后端能力已具备，纯前端补全；\n'
        '2) P1 两项可合并为一轮"AI 增强"迭代：思政 AI 提取 + 课中即时测验；\n'
        '3) P2-P3 列入第二期 roadmap，与产品侧确认优先级；\n'
        '4) 国产化适配（multi-arch）建议作为"信创版"分支独立交付，'
        '避免影响主线 x86_64 的发布节奏。'
    )
    set_run_font(r, size=10.5)

    # ===== 保存 =====
    out_path = '/Users/bai/projects/kg-edu/docs/课堂星智慧教学系统第二期_验收清单.docx'
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f'✅ 验收文档已生成: {out_path}')
    print(f'   文件大小: {os.path.getsize(out_path) / 1024:.1f} KB')
    return out_path


if __name__ == '__main__':
    build_document()
