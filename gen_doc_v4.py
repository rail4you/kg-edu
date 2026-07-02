#!/usr/bin/env python3
"""智课云枢 - 功能参数说明书（全量版）"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = '\u5fae\u8f6f\u96c5\u9ed1'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

def shade(cell, c):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{c}"/>'))
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '\u5fae\u8f6f\u96c5\u9ed1'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
def para(text, bold=False, size=10.5, sa=4, sb=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    run = p.add_run(text)
    run.font.name = '\u5fae\u8f6f\u96c5\u9ed1'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
def make_tbl(headers, data, cw=None):
    t = doc.add_table(rows=0, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for is_hdr, row_data in [(True, headers)] + [(False, d) for d in data]:
        row = t.add_row()
        for i, txt in enumerate(row_data):
            cell = row.cells[i]; cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run(str(txt))
            run.font.name = '\u5fae\u8f6f\u96c5\u9ed1'
            run.font.size = Pt(9.5)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
            if is_hdr:
                run.bold = True
                shade(cell, 'D9E2F3')
    if cw:
        for i, w in enumerate(cw):
            for row in t.rows:
                row.cells[i].width = Cm(w)
def param_table(title, items):
    tbl = doc.add_table(rows=0, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    for is_hdr, row_data in [(True, ['\u5e8f\u53f7', '\u529f\u80fd\u9879', '\u6280\u672f\u53c2\u6570\u8981\u6c42'])] + [(False, it) for it in items]:
        row = tbl.add_row()
        for i, txt in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            run = p.add_run(str(txt))
            run.font.name = '\u5fae\u8f6f\u96c5\u9ed1'
            run.font.size = Pt(9)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
            if is_hdr:
                run.bold = True
                shade(cell, 'D9E2F3')
    for i, w in enumerate([1.5, 3.2, 12.5]):
        for row in tbl.rows:
            row.cells[i].width = Cm(w)

# ====== 封面 ======
for _ in range(2):
    doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tp.add_run('\u667a\u8bfe\u4e91\u67a2 \u667a\u6167\u6559\u5b66\u7cfb\u7edf')
run.font.name = '\u5fae\u8f6f\u96c5\u9ed1'
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sp.add_run('\u529f\u80fd\u53c2\u6570\u8bf4\u660e\u4e66\uff08\u5168\u91cf\u7248\uff09')
run.font.name = '\u5fae\u8f6f\u96c5\u9ed1'
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
doc.add_paragraph()
para('\u4e3a\u79ef\u6781\u843d\u5b9e\u56fd\u5bb6\u6559\u80b2\u6570\u5b57\u5316\u6218\u7565\u884c\u52a8\uff0c\u63a8\u52a8\u4eba\u5de5\u667a\u80fd\u8d4b\u80fd\u6559\u80b2\u53d8\u9769\u8f6c\u578b\uff0c\u5b9e\u73b0\u4fe1\u606f\u6280\u672f\u4e0e\u6559\u80b2\u6559\u5b66\u6df1\u5ea6\u878d\u5408\uff0c\u63d0\u9ad8\u8bfe\u7a0b\u5efa\u8bbe\u8d28\u91cf\uff0c\u52a9\u529b\u4e13\u4e1a\u6570\u667a\u5316\u8f6c\u578b\uff0c\u63d0\u5347\u4eba\u624d\u57f9\u517b\u6210\u6548\uff0c\u62df\u5bf9\u73b0\u6709\u5728\u7ebf\u7cbe\u54c1\u8bfe\u7a0b\u8fdb\u884c\u201c\u4eba\u5de5\u667a\u80fd\uff0b\u201d\u667a\u80fd\u5347\u7ea7\uff0c\u6253\u9020\u201c\u667a\u8bfe\u4e91\u67a2\u201d\u667a\u6167\u6559\u5b66\u7cfb\u7edf\u3002', size=11, sa=6)
para('\u201c\u667a\u8bfe\u4e91\u67a2\u201d\u5229\u7528\u5927\u6a21\u578b\u3001\u77e5\u8bc6\u56fe\u8c31\u7b49\u4eba\u5de5\u667a\u80fd\u6280\u672f\u4e0e\u804c\u4e1a\u6559\u80b2/\u9ad8\u7b49\u6559\u80b2\u6df1\u5ea6\u878d\u5408\uff0c\u4e3a\u5b66\u6821\u6253\u9020\u667a\u80fd\u5316\u3001\u521b\u65b0\u6027\u3001\u5b9e\u7528\u6027\u7684\u667a\u6167\u8bfe\u7a0b\u5e73\u53f0\u3002', size=11, sa=6)
para('\u6280\u672f\u6808\uff1a\u524d\u7aef React 18 + TypeScript + Vite + Ant Design Pro\uff08\u8682\u8681\u96c6\u56e2\uff09\uff1b\u540e\u7aef Elixir + Phoenix + Ash Framework\uff1b\u6570\u636e\u5e93 PostgreSQL\uff08\u591a\u79df\u6237 Schema \u9694\u79bb\uff09\uff1b\u5bf9\u8c61\u5b58\u50a8\u963f\u91cc\u4e91 OSS\uff1bAI \u5f15\u64ce\u57fa\u4e8e\u901a\u4e49\u5343\u95ee\uff08\u963f\u91cc\u4e91\u767e\u70bc\uff09\u3002', size=10.5, sa=4)
doc.add_paragraph()
vp = doc.add_paragraph()
vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = vp.add_run('\u7248\u672c\uff1aV2.0\uff08\u5168\u91cf\u7248\uff09 | 2025\u5e746\u6708')
run.font.name = '\u5fae\u8f6f\u96c5\u9ed1'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
doc.add_page_break()

print('Script loaded OK, generating document...')

# ====== Write document sections ======
import json

sections_data = json.loads(open(os.path.join(os.path.dirname(__file__) or '.', 'sections_data.json')).read())

for sec_title, sec_desc, items in sections_data:
    doc.add_page_break()
    heading(sec_title, 1)
    para(sec_desc)
    formatted = [[it[0], it[1], it[2].replace('|', '\n')] for it in items]
    param_table(sec_title, formatted)

# Appendix
doc.add_page_break()
heading('\u9644\u5f55\uff1a\u6280\u672f\u53c2\u6570\u6c47\u603b\u8868', 1)
appendix_data = json.loads(open(os.path.join(os.path.dirname(__file__) or '.', 'appendix_data.json')).read())
make_tbl(['\u6280\u672f\u6307\u6807', '\u53c2\u6570/\u8bf4\u660e'], appendix_data, [4.5, 12])

out = os.path.expanduser('~/projects/kg-edu/\u667a\u8bfe\u4e91\u67a2_\u529f\u80fd\u53c2\u6570\u8bf4\u660e\u4e66\uff08\u5168\u91cf\u7248\uff09.docx')
doc.save(out)
print(f'\u2705 \u6587\u6863\u5df2\u751f\u6210: {out}')
